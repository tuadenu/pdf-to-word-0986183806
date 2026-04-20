import threading
import traceback
import subprocess
import time
from dataclasses import dataclass
from pathlib import Path
from queue import Queue
from typing import Optional

import fitz  # PyMuPDF
import pytesseract
import tkinter as tk
from pdf2docx import Converter
from PIL import Image
from tkinter import filedialog, messagebox, ttk
from docx import Document


@dataclass
class ConvertResult:
    file_name: str
    success: bool
    message: str


def has_embedded_text(pdf_path: Path, max_pages_to_check: int = 3, min_chars: int = 30) -> bool:
    """Quickly detect whether PDF likely contains selectable text."""
    try:
        doc = fitz.open(pdf_path)
        total = min(len(doc), max_pages_to_check)
        chars = 0
        for i in range(total):
            chars += len(doc[i].get_text("text").strip())
            if chars >= min_chars:
                doc.close()
                return True
        doc.close()
    except Exception:
        return False
    return False


def ocr_pdf_to_docx(pdf_path: Path, output_docx_path: Path, lang: str) -> None:
    """OCR each PDF page and write recognized text to a DOCX file."""
    doc = fitz.open(pdf_path)
    out_doc = Document()

    for idx, page in enumerate(doc):
        # Render page at higher resolution to improve OCR quality.
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

        text = pytesseract.image_to_string(img, lang=lang)
        clean_text = text.strip()

        out_doc.add_heading(f"Page {idx + 1}", level=2)
        if clean_text:
            for paragraph in clean_text.split("\n"):
                paragraph = paragraph.strip()
                if paragraph:
                    out_doc.add_paragraph(paragraph)
        else:
            out_doc.add_paragraph("[No text recognized on this page]")

    doc.close()
    out_doc.save(output_docx_path)


def convert_one_pdf(
    pdf_path: Path,
    output_dir: Path,
    use_ocr_fallback: bool,
    ocr_lang: str,
    preserve_layout: bool = True,
) -> ConvertResult:
    output_path = output_dir / f"{pdf_path.stem}.docx"

    try:
        if has_embedded_text(pdf_path):
            converter = Converter(str(pdf_path))
            converter.convert(str(output_path), start=0, end=None, preserve_layout=preserve_layout)
            converter.close()
            return ConvertResult(pdf_path.name, True, "Trich xuat tu text co san")

        if use_ocr_fallback:
            ocr_pdf_to_docx(pdf_path, output_path, ocr_lang)
            return ConvertResult(pdf_path.name, True, "Converted with OCR fallback")

        return ConvertResult(
            pdf_path.name,
            False,
            "No embedded text found. Enable OCR fallback to convert scanned PDFs.",
        )

    except Exception as exc:
        return ConvertResult(pdf_path.name, False, f"{exc}\n{traceback.format_exc(limit=1)}")


class App:
    def __init__(self, root: tk.Tk) -> None:
        self.root = root
        self.root.title("Chuyen PDF sang Word")
        self.root.geometry("1080x700")
        self.root.minsize(980, 640)

        self.input_dir_var = tk.StringVar()
        self.output_dir_var = tk.StringVar()
        self.ocr_var = tk.BooleanVar(value=True)
        self.ocr_lang_var = tk.StringVar(value="eng+vie+chi_sim+jpn")
        self.skip_existing_var = tk.BooleanVar(value=False)
        self.preserve_layout_var = tk.BooleanVar(value=True)
        self.use_file_mode_var = tk.BooleanVar(value=False)
        self.status_var = tk.StringVar(value="San sang")
        self.selected_files: list[Path] = []
        self.summary_var = tk.StringVar(value="Chua chon tep")
        self.progress_var = tk.DoubleVar(value=0)
        self.total_files = 0
        self.processed_files = 0
        self.success_count = 0
        self.is_running = False
        self.stop_requested = False
        self.current_thread: Optional[threading.Thread] = None
        self.row_map: dict[str, str] = {}
        self.start_time: Optional[float] = None
        self.avg_time_per_file = 0.0

        self.log_queue: Queue[str] = Queue()
        self._setup_theme()
        self._build_ui()
        self._poll_log_queue()

    def _setup_theme(self) -> None:
        style = ttk.Style()
        available = set(style.theme_names())
        if "clam" in available:
            style.theme_use("clam")

        style.configure("App.TFrame", background="#f4f7fb")
        style.configure("Card.TFrame", background="#ffffff")
        style.configure("Title.TLabel", font=("Avenir Next", 20, "bold"), background="#f4f7fb", foreground="#12243a")
        style.configure("Sub.TLabel", font=("Avenir Next", 10), background="#f4f7fb", foreground="#425466")
        style.configure("CardTitle.TLabel", font=("Avenir Next", 12, "bold"), background="#ffffff", foreground="#18314f")
        style.configure("Status.TLabel", font=("Avenir Next", 10, "bold"), background="#ffffff", foreground="#1f4a7c")
        style.configure("Primary.TButton", font=("Avenir Next", 10, "bold"))

        style.configure(
            "App.Horizontal.TProgressbar",
            troughcolor="#dce6f2",
            background="#2f7ed8",
            bordercolor="#dce6f2",
            lightcolor="#2f7ed8",
            darkcolor="#2f7ed8",
        )

    def _build_ui(self) -> None:
        self.root.configure(bg="#f4f7fb")
        container = ttk.Frame(self.root, style="App.TFrame", padding=14)
        container.pack(fill="both", expand=True)

        header = ttk.Frame(container, style="App.TFrame")
        header.pack(fill="x", pady=(0, 12))

        ttk.Label(
            header,
            text="Chuyen PDF sang Word",
            style="Title.TLabel",
        ).pack(anchor="w")
        ttk.Label(
            header,
            text="Chuyen doi hang loat sang DOCX co the chinh sua, co OCR cho file scan",
            style="Sub.TLabel",
        ).pack(anchor="w", pady=(2, 0))

        settings_card = ttk.Frame(container, style="Card.TFrame", padding=14)
        settings_card.pack(fill="x", pady=(0, 12))

        ttk.Label(settings_card, text="Cai dat", style="CardTitle.TLabel").grid(row=0, column=0, sticky="w", pady=(0, 10))

        # Mode selection: Directory or Files
        mode_frame = ttk.Frame(settings_card, style="Card.TFrame")
        mode_frame.grid(row=1, column=0, columnspan=3, sticky="w", pady=(0, 12))
        self.mode_dir_radio = ttk.Radiobutton(mode_frame, text="Chuyen toan bo thu muc", variable=self.use_file_mode_var, value=False, command=self._on_mode_changed)
        self.mode_dir_radio.pack(side="left", padx=(0, 12))
        self.mode_file_radio = ttk.Radiobutton(mode_frame, text="Chon file le (Ctrl/Cmd + Click)", variable=self.use_file_mode_var, value=True, command=self._on_mode_changed)
        self.mode_file_radio.pack(side="left")

        ttk.Label(settings_card, text="Thu muc dau vao (PDF):", width=20).grid(row=2, column=0, sticky="w", padx=(0, 8), pady=4)
        self.input_entry = ttk.Entry(settings_card, textvariable=self.input_dir_var)
        self.input_entry.grid(row=2, column=1, sticky="ew", pady=4)
        self.input_button = ttk.Button(settings_card, text="Chon", command=self._choose_input_dir)
        self.input_button.grid(row=2, column=2, sticky="ew", padx=(8, 0), pady=4)

        # File selection section
        self.file_select_frame = ttk.Frame(settings_card, style="Card.TFrame")
        self.file_select_frame.grid(row=3, column=0, columnspan=3, sticky="ew", pady=(8, 0))
        self.file_select_button = ttk.Button(self.file_select_frame, text="Chon file PDF", command=self._choose_pdf_files)
        self.file_select_button.pack(side="left", padx=(0, 8))
        self.file_clear_button = ttk.Button(self.file_select_frame, text="Xoa danh sach", command=self._clear_file_selection)
        self.file_clear_button.pack(side="left", padx=(0, 8))
        self.file_count_label = ttk.Label(self.file_select_frame, text="0 file", style="Sub.TLabel")
        self.file_count_label.pack(side="left")
        self.file_select_frame.grid_remove()  # Hide initially

        ttk.Label(settings_card, text="Thu muc dau ra (DOCX):", width=20).grid(row=4, column=0, sticky="w", padx=(0, 8), pady=4)
        self.output_entry = ttk.Entry(settings_card, textvariable=self.output_dir_var)
        self.output_entry.grid(row=4, column=1, sticky="ew", pady=4)
        self.output_button = ttk.Button(settings_card, text="Chon", command=self._choose_output_dir)
        self.output_button.grid(row=4, column=2, sticky="ew", padx=(8, 0), pady=4)

        self.ocr_checkbox = ttk.Checkbutton(
            settings_card,
            text="Bat OCR du phong cho PDF scan",
            variable=self.ocr_var,
        )
        self.ocr_checkbox.grid(row=5, column=0, columnspan=2, sticky="w", pady=(8, 4))

        self.skip_checkbox = ttk.Checkbutton(
            settings_card,
            text="Bo qua file DOCX da ton tai",
            variable=self.skip_existing_var,
        )
        self.skip_checkbox.grid(row=6, column=0, columnspan=2, sticky="w", pady=(2, 8))

        self.preserve_checkbox = ttk.Checkbutton(
            settings_card,
            text="Giu dinh dang (font, trang, bo cuc)",
            variable=self.preserve_layout_var,
        )
        self.preserve_checkbox.grid(row=7, column=0, columnspan=2, sticky="w", pady=(2, 8))

        lang_row = ttk.Frame(settings_card, style="Card.TFrame")
        lang_row.grid(row=8, column=0, columnspan=3, sticky="w", pady=(2, 0))
        ttk.Label(lang_row, text="Ngon ngu OCR:").pack(side="left")
        self.lang_entry = ttk.Entry(lang_row, textvariable=self.ocr_lang_var, width=12)
        self.lang_entry.pack(side="left", padx=(6, 8))
        ttk.Label(lang_row, text="Vi du: eng, vie, chi_sim, jpn, kor (hoac chi_sim+jpn)", style="Sub.TLabel").pack(side="left")

        settings_card.columnconfigure(1, weight=1)

        controls_card = ttk.Frame(container, style="Card.TFrame", padding=14)
        controls_card.pack(fill="x", pady=(0, 12))

        ttk.Label(controls_card, text="Tien do", style="CardTitle.TLabel").grid(row=0, column=0, sticky="w")
        ttk.Label(controls_card, textvariable=self.status_var, style="Status.TLabel").grid(row=0, column=1, sticky="e")

        self.progress_bar = ttk.Progressbar(
            controls_card,
            variable=self.progress_var,
            maximum=100,
            style="App.Horizontal.TProgressbar",
        )
        self.progress_bar.grid(row=1, column=0, columnspan=2, sticky="ew", pady=(8, 4))

        ttk.Label(controls_card, textvariable=self.summary_var).grid(row=2, column=0, sticky="w")

        actions = ttk.Frame(controls_card, style="Card.TFrame")
        actions.grid(row=2, column=1, sticky="e")
        self.start_button = ttk.Button(actions, text="Bat dau", style="Primary.TButton", command=self._start_conversion)
        self.start_button.pack(side="left")
        self.stop_button = ttk.Button(actions, text="Dung", command=self._request_stop, state="disabled")
        self.stop_button.pack(side="left", padx=(8, 0))
        self.open_output_button = ttk.Button(actions, text="Mo thu muc", command=self._open_output_folder)
        self.open_output_button.pack(side="left", padx=(8, 0))

        controls_card.columnconfigure(0, weight=1)

        table_card = ttk.Frame(container, style="Card.TFrame", padding=14)
        table_card.pack(fill="both", expand=True)

        ttk.Label(table_card, text="Danh sach tep", style="CardTitle.TLabel").pack(anchor="w", pady=(0, 8))

        columns = ("file", "status", "detail")
        self.results_tree = ttk.Treeview(table_card, columns=columns, show="headings", height=14)
        self.results_tree.heading("file", text="Tep")
        self.results_tree.heading("status", text="Trang thai")
        self.results_tree.heading("detail", text="Chi tiet")
        self.results_tree.column("file", width=280, anchor="w")
        self.results_tree.column("status", width=100, anchor="center")
        self.results_tree.column("detail", width=520, anchor="w")

        scrollbar = ttk.Scrollbar(table_card, orient="vertical", command=self.results_tree.yview)
        self.results_tree.configure(yscrollcommand=scrollbar.set)
        self.results_tree.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        self.results_tree.tag_configure("queued", foreground="#5f6b7a")
        self.results_tree.tag_configure("ok", foreground="#0b6a33")
        self.results_tree.tag_configure("fail", foreground="#a11e1e")

    def _on_mode_changed(self) -> None:
        if self.use_file_mode_var.get():
            self.input_entry.configure(state="disabled")
            self.input_button.configure(state="disabled")
            self.file_select_frame.grid()
        else:
            self.input_entry.configure(state="normal")
            self.input_button.configure(state="normal")
            self.file_select_frame.grid_remove()
            self._clear_file_selection()

    def _choose_input_dir(self) -> None:
        path = filedialog.askdirectory(title="Chon thu muc chua file PDF")
        if path:
            self.input_dir_var.set(path)

    def _choose_pdf_files(self) -> None:
        files = filedialog.askopenfilenames(
            title="Chon file PDF (Ctrl/Cmd + Click de chon nhieu file)",
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")],
        )
        if files:
            self.selected_files = [Path(f) for f in files]
            self._update_file_count()

    def _clear_file_selection(self) -> None:
        self.selected_files.clear()
        self._update_file_count()

    def _update_file_count(self) -> None:
        count = len(self.selected_files)
        self.file_count_label.configure(text=f"{count} file" if count != 1 else "1 file")

    def _choose_output_dir(self) -> None:
        path = filedialog.askdirectory(title="Chon thu muc luu file DOCX")
        if path:
            self.output_dir_var.set(path)

    def _append_log(self, text: str) -> None:
        self.status_var.set(text)

    def _poll_log_queue(self) -> None:
        while not self.log_queue.empty():
            message = self.log_queue.get_nowait()
            self._append_log(message)
        self.root.after(120, self._poll_log_queue)

    def _log(self, text: str) -> None:
        self.log_queue.put(text)

    def _validate_paths(self) -> Optional[tuple[Path, Path]]:
        input_dir = Path(self.input_dir_var.get().strip())
        output_dir = Path(self.output_dir_var.get().strip())

        if not input_dir.exists() or not input_dir.is_dir():
            messagebox.showerror("Dau vao khong hop le", "Vui long chon thu muc dau vao hop le.")
            return None

        if not output_dir.exists():
            output_dir.mkdir(parents=True, exist_ok=True)

        if not output_dir.is_dir():
            messagebox.showerror("Dau ra khong hop le", "Vui long chon thu muc dau ra hop le.")
            return None

        return input_dir, output_dir

    def _set_controls_enabled(self, enabled: bool) -> None:
        state = "normal" if enabled else "disabled"
        self.start_button.configure(state=state)
        self.mode_dir_radio.configure(state=state)
        self.mode_file_radio.configure(state=state)
        self.output_entry.configure(state=state)
        self.output_button.configure(state=state)
        self.ocr_checkbox.configure(state=state)
        self.skip_checkbox.configure(state=state)
        self.preserve_checkbox.configure(state=state)
        self.lang_entry.configure(state=state)
        if self.use_file_mode_var.get():
            self.file_select_button.configure(state=state)
            self.file_clear_button.configure(state=state)
        else:
            self.input_entry.configure(state=state)
            self.input_button.configure(state=state)
        self.stop_button.configure(state="disabled" if enabled else "normal")
        self.is_running = not enabled

    def _open_output_folder(self) -> None:
        output_dir = Path(self.output_dir_var.get().strip())
        if not output_dir.exists() or not output_dir.is_dir():
            messagebox.showwarning("Mo thu muc", "Vui long chon thu muc dau ra hop le truoc.")
            return

        try:
            self.root.tk.call("::tk::mac::OpenDocument", str(output_dir))
        except tk.TclError:
            # Fallback for platforms where macOS helper is unavailable.
            subprocess.run(["open", str(output_dir)], check=False)

    def _reset_result_table(self, pdf_files: list[Path]) -> None:
        for item_id in self.results_tree.get_children():
            self.results_tree.delete(item_id)

        self.row_map.clear()
        for pdf in pdf_files:
            row_id = self.results_tree.insert(
                "",
                "end",
                values=(pdf.name, "Cho xu ly", "Dang doi"),
                tags=("queued",),
            )
            self.row_map[pdf.name] = row_id

    def _update_row(self, file_name: str, status: str, detail: str, ok: bool) -> None:
        row_id = self.row_map.get(file_name)
        if not row_id:
            return
        tag = "ok" if ok else "fail"
        self.results_tree.item(row_id, values=(file_name, status, detail), tags=(tag,))

    def _update_progress(self) -> None:
        if self.total_files <= 0:
            self.progress_var.set(0)
            self.summary_var.set("Chua chon tep")
            return

        percent = (self.processed_files / self.total_files) * 100
        self.progress_var.set(percent)

        # Calculate estimated time remaining
        time_info = ""
        if self.start_time and self.processed_files > 0:
            elapsed = time.time() - self.start_time
            self.avg_time_per_file = elapsed / self.processed_files
            remaining_files = self.total_files - self.processed_files
            est_remaining_sec = int(remaining_files * self.avg_time_per_file)

            if est_remaining_sec > 0:
                minutes = est_remaining_sec // 60
                seconds = est_remaining_sec % 60
                if minutes > 0:
                    time_info = f" | Con ~{minutes}m {seconds}s"
                else:
                    time_info = f" | Con ~{seconds}s"

        self.summary_var.set(
            f"Da xu ly {self.processed_files}/{self.total_files} | Thanh cong: {self.success_count}{time_info}"
        )

    def _on_file_done(self, result: ConvertResult) -> None:
        if result.success:
            self.success_count += 1
            self._update_row(result.file_name, "Thanh cong", result.message, ok=True)
        else:
            self._update_row(result.file_name, "That bai", result.message.split("\n")[0], ok=False)

        self.processed_files += 1
        self._update_progress()

    def _on_batch_done(self, stopped: bool) -> None:
        if stopped:
            done_text = (
                f"Da dung: {self.success_count}/{self.total_files} tep thanh cong, "
                f"da xu ly {self.processed_files} tep"
            )
            title = "Da dung"
        else:
            done_text = f"Hoan tat: {self.success_count}/{self.total_files} tep thanh cong"
            title = "Hoan tat"

        self._log(done_text)
        self._set_controls_enabled(True)
        messagebox.showinfo(title, done_text)

    def _request_stop(self) -> None:
        if not self.is_running:
            return
        self.stop_requested = True
        self.status_var.set("Dang dung... se dung sau khi xong tep hien tai")

    def _start_conversion(self) -> None:
        if self.is_running:
            return

        output_dir = Path(self.output_dir_var.get().strip())
        if not output_dir.exists():
            output_dir.mkdir(parents=True, exist_ok=True)
        if not output_dir.is_dir():
            messagebox.showerror("Dau ra khong hop le", "Vui long chon thu muc dau ra hop le.")
            return

        # Get PDF files based on mode
        if self.use_file_mode_var.get():
            if not self.selected_files:
                messagebox.showinfo("Khong co tep", "Vui long chon it nhat 1 file PDF.")
                return
            pdf_files = sorted(self.selected_files)
        else:
            paths = self._validate_paths()
            if not paths:
                return
            input_dir, _ = paths
            pdf_files = sorted([p for p in input_dir.iterdir() if p.is_file() and p.suffix.lower() == ".pdf"])

        # Filter out files if skip_existing is enabled
        if self.skip_existing_var.get():
            pdf_files = [p for p in pdf_files if not (output_dir / f"{p.stem}.docx").exists()]

        if not pdf_files:
            if self.use_file_mode_var.get():
                messagebox.showinfo("Khong co tep", "Tat ca file da duoc chuyen doi hoac bi bo qua.")
            else:
                messagebox.showinfo("Khong co tep", "Khong tim thay file PDF trong thu muc dau vao.")
            self.status_var.set("San sang")
            self.progress_var.set(0)
            self.summary_var.set("Chua chon tep")
            return

        self.stop_requested = False
        self.start_time = time.time()
        self.total_files = len(pdf_files)
        self.processed_files = 0
        self.success_count = 0
        self.avg_time_per_file = 0.0
        self.progress_var.set(0)
        self._reset_result_table(pdf_files)
        self._update_progress()
        self._set_controls_enabled(False)

        worker = threading.Thread(
            target=self._run_batch_conversion,
            args=(output_dir, output_dir, pdf_files, self.ocr_var.get(), self.ocr_lang_var.get().strip() or "eng", self.preserve_layout_var.get()),
            daemon=True,
        )
        self.current_thread = worker
        worker.start()

    def _run_batch_conversion(
        self,
        input_dir: Path,
        output_dir: Path,
        pdf_files: list[Path],
        use_ocr_fallback: bool,
        ocr_lang: str,
        preserve_layout: bool,
    ) -> None:
        self._log("Dang chuyen doi...")
        self._log(f"Dau vao: {input_dir}")

        stopped = False
        remaining_files: list[Path] = []

        for idx, pdf in enumerate(pdf_files):
            if self.stop_requested:
                stopped = True
                remaining_files = pdf_files[idx:]
                break

            result = convert_one_pdf(
                pdf,
                output_dir,
                use_ocr_fallback=use_ocr_fallback,
                ocr_lang=ocr_lang,
                preserve_layout=preserve_layout,
            )

            if result.success:
                if "embedded text extraction" in result.message:
                    result.message = "Trich xuat tu text co san"
                elif "OCR fallback" in result.message:
                    result.message = "Nhan dang bang OCR"
            else:
                result.message = result.message.replace(
                    "No embedded text found. Enable OCR fallback to convert scanned PDFs.",
                    "Khong co text co san. Hay bat OCR de chuyen doi PDF scan.",
                )

            self.root.after(0, self._on_file_done, result)

        if stopped:
            for pdf in remaining_files:
                self.root.after(0, self._update_row, pdf.name, "Da huy", "Nguoi dung da dung", False)

        self.root.after(0, self._on_batch_done, stopped)


def main() -> None:
    root = tk.Tk()
    app = App(root)
    root.mainloop()


if __name__ == "__main__":
    main()
